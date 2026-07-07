import datetime as dt
from pprint import pprint
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Mm, Pt
import xlsxwriter

from .models import Statement, Productions


def get_working_minutes(product: Statement):
    try:
        if product.statement_end_time and product.statement_start_time:
            start_time = dt.datetime.combine(product.statement_date, product.statement_start_time)
            end_time = dt.datetime.combine(product.statement_date, product.statement_end_time)
            return int((end_time - start_time).total_seconds() // 60)
        else:
            return 0
    except:
        return 0


def defects_percent(product: Statement):
    try:
        if product.quantity_sent_production and product.quantity_manufactured:
            k_defects = (product.quantity_sent_production - product.quantity_manufactured) / product.quantity_manufactured
            result = round(k_defects * 100.0, 1)
            return result
        else:
            return  "-"
    except:
        return "-"


def summ_defects_percent(quantity_sent_production: int, quantity_manufactured: int):
    try:
        if quantity_manufactured != 0:
            k_defects = (quantity_sent_production - quantity_manufactured) / quantity_manufactured
            result = round(k_defects * 100.0, 1)
            return result
        else:
            return  "-"
    except:
        return "-"


def speed_manufactured(product: Statement):
    try:
        if product.quantity_manufactured:
            return round(product.quantity_manufactured / get_working_minutes(product), 1)
        else:
            return  "-"
    except:
        return "-"


def summ_speed_manufactured(quantity_manufactured:int, summ_working_minutes:int):
    try:
        if summ_working_minutes != 0:
            return round(quantity_manufactured / summ_working_minutes, 1)
        else:
            return  "-"
    except:
        return "-"


def manufactured_area(product: Statement):
    try:
        if product.order and product.quantity_manufactured:
            area = product.order.width * product.order.length * product.quantity_manufactured / (1000 * 1000)
            return round(area, 1)
        return 0
    except:
        return 0


def statement_to_excel(workbook:xlsxwriter.Workbook, statement_date:str):
    ws = workbook.add_worksheet('Ведомость')

    ws.set_row(4, 50)

    ws.set_column(0, 0, 10)
    ws.set_column(1, 1, 40)
    ws.set_column(2, 3, 10)
    ws.set_column(4, 6, 12)
    ws.set_column(7, 9, 10)
    ws.set_column(10, 11, 13)
    ws.set_column(12, 16, 12)

    title_format = workbook.add_format({'bold': True, 'align': 'center'})

    bold_format = workbook.add_format({'bold': True, 'border': 2, 'align': 'center'})
    bold_format.set_align('vcenter')
    bold_format.set_text_wrap(True)
    left_format = workbook.add_format({'bold': False, 'border': 1})
    def_format = workbook.add_format({'bold': False, 'border': 1, 'align': 'center'})

    ws.merge_range("A3:P3", 'Ведомость  данных изготовления  продукции на линиях', title_format)

    row_offset = 4
    ws.write(row_offset, 0, 'Дата', bold_format)
    ws.write(row_offset, 1, 'Наименование продукции', bold_format)
    ws.write(row_offset, 2, 'Печать', bold_format)
    ws.write(row_offset, 3, 'Штамп', bold_format)
    ws.write(row_offset, 4, 'Ширина, мм', bold_format)
    ws.write(row_offset, 5, 'Длина, мм', bold_format)
    ws.write(row_offset, 6, 'Площадь м²', bold_format)
    ws.write(row_offset, 7, 'Начало', bold_format)
    ws.write(row_offset, 8, 'Окончан.', bold_format)
    ws.write(row_offset, 9, 'Минуты', bold_format)
    ws.write(row_offset, 10, 'Пропущено шт', bold_format)
    ws.write(row_offset, 11, 'Изготовлено шт', bold_format)
    ws.write(row_offset, 12, 'Брак %', bold_format)
    ws.write(row_offset, 13, 'Произвед. шт/мин', bold_format)
    ws.write(row_offset, 14, 'Простой', bold_format)
    ws.write(row_offset, 15, 'Общая площадь м²', bold_format)

    summ_working_minutes = 0
    summ_quantity_sent_production = 0
    summ_quantity_manufactured = 0
    summ_manufactured_area = 0

    queryset = Statement.objects.filter(statement_date=statement_date).order_by('statement_start_time')
    # print(queryset)
    for product in queryset:
        row_offset += 1
        ws.write(row_offset, 0, product.statement_date.strftime('%d.%m.%Y'), def_format)
        ws.write(row_offset, 1, product.order.name, left_format)
        color_count = str(product.order.color_count) if product.order.color_count != 0 else '-'
        ws.write(row_offset, 2, color_count, def_format)
        stamp = '+' if product.order.stamp else '-'
        ws.write(row_offset, 3, stamp, def_format)
        ws.write(row_offset, 4, product.order.width, def_format)
        ws.write(row_offset, 5, product.order.length, def_format)
        ws.write(row_offset, 6, product.order.area, def_format)
        ws.write(row_offset, 7, product.statement_start_time.strftime('%H:%M'), def_format)
        ws.write(row_offset, 8, product.statement_end_time.strftime('%H:%M'), def_format)
        summ_working_minutes += get_working_minutes(product)
        ws.write(row_offset, 9, get_working_minutes(product), def_format)
        summ_quantity_sent_production += product.quantity_sent_production
        ws.write(row_offset, 10, product.quantity_sent_production, def_format)
        summ_quantity_manufactured += product.quantity_manufactured
        ws.write(row_offset, 11, product.quantity_manufactured, def_format)
        ws.write(row_offset, 12, defects_percent(product), def_format)
        ws.write(row_offset, 13, speed_manufactured(product), def_format)
        ws.write(row_offset, 14, '', def_format)
        summ_manufactured_area += manufactured_area(product)
        ws.write(row_offset, 15, manufactured_area(product), def_format)

    row_offset += 1
    ws.write(row_offset, 0, 'Итого', bold_format)
    ws.write(row_offset, 1, '', bold_format)
    ws.write(row_offset, 2, '', bold_format)
    ws.write(row_offset, 3, '', bold_format)
    ws.write(row_offset, 4, '', bold_format)
    ws.write(row_offset, 5, '', bold_format)
    ws.write(row_offset, 6, '', bold_format)
    ws.write(row_offset, 7, '', bold_format)
    ws.write(row_offset, 8, '', bold_format)
    ws.write(row_offset, 9, summ_working_minutes, bold_format)
    ws.write(row_offset, 10, summ_quantity_sent_production, bold_format)
    ws.write(row_offset, 11, summ_quantity_manufactured, bold_format)
    ws.write(row_offset, 12, summ_defects_percent(quantity_sent_production=summ_quantity_sent_production,
                                                            quantity_manufactured=summ_quantity_manufactured), bold_format)
    ws.write(row_offset, 13, summ_speed_manufactured(quantity_manufactured=summ_quantity_manufactured,
                                                               summ_working_minutes=summ_working_minutes), bold_format)
    ws.write(row_offset, 14, '', bold_format)
    ws.write(row_offset, 15, summ_manufactured_area, bold_format)

    ws.print_area(0, 0, 15, 15)
    ws.set_paper(9)
    ws.set_landscape()
    ws.fit_to_pages(1, 1)


def invoice_to_word(doc, invoice_date:str):
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)

    section.left_margin = Mm(10)
    section.right_margin = Mm(154)
    section.top_margin = Mm(5)
    section.bottom_margin = Mm(5)
    #
    # par0.paragraph_format.space_before = Mm(1)
    # par0.paragraph_format.space_after = Mm(1)

    # Данные для подстановки
    par0 = doc.add_paragraph(f'{invoice_date} г.')
    par1 = doc.add_paragraph(f'Накладная №_______________')
    par1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par2 = doc.add_paragraph(f'От цеха по изготовлению упаковки')
    par2.add_run(f'\nВ склад готовой продукции')

    # Создание таблицы для позиций
    table = doc.add_table(rows=1, cols=5)
    table.autofit = False
    table.allow_autofit = False
    table.style = 'Table Grid'
    # table.width = Mm(131.3)

    table.columns[0].width = Mm(20)
    table.columns[1].width = Mm(10)
    table.columns[2].width = Mm(70)
    table.columns[3].width = Mm(20)
    table.columns[4].width = Mm(20)

    table.cell(0, 0).text = 'S'
    table.cell(0, 1).text = '№'
    table.cell(0, 2).text = 'Наименование \nпродукции'
    table.cell(0, 3).text = 'Печать'
    table.cell(0, 4).text = 'Количество'
    table.cell(0, 2).bold = True  # Выделяем жирным название колонки

    # Заполнение таблицы данными
    products = Statement.objects.filter(statement_date=invoice_date).order_by('statement_start_time')
    n = 0
    for product in products:
        if product.order_status == Productions.Status.MANUFACTURED:
            n += 1
            row_cells = table.add_row().cells
            row_cells[0].text = str(manufactured_area(product))
            row_cells[1].text = str(n)
            row_cells[2].text = str(product.order.name)
            row_cells[3].text = str('+' if product.order.stamp else '-')
            row_cells[4].text = str(product.quantity_manufactured)

    for _ in range(21 - n):
        table.add_row()

    par_end = doc.add_paragraph(f'Сдал_____________           Принял_____________')
    par_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
